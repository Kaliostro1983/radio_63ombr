from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, date, time
from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .enemy_moves_config import load_reports_config


@dataclass(frozen=True)
class EnemyMoveItem:
    freq: str
    move: str
    unit: str
    group: str


def _set_base_font_12(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.size = Pt(12)


def _add_header(doc: Document, *, title: str, subtitle: str) -> None:
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(title)
    run1.bold = True
    run1.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(14)

    doc.add_paragraph("")


def _group_items(items: Iterable[EnemyMoveItem]) -> dict[str, list[EnemyMoveItem]]:
    grouped: dict[str, list[EnemyMoveItem]] = {}
    for it in items:
        key = (it.group or "").strip()
        if not key:
            continue
        grouped.setdefault(key, []).append(it)
    return grouped


def build_enemy_moves_docx_bytes(items: list[EnemyMoveItem], *, report_date: date | None = None) -> tuple[bytes, str]:
    cfg = load_reports_config().get("enemy_moves") or {}
    if not isinstance(cfg, dict):
        cfg = {}

    today = report_date or datetime.now().date()
    date_str = today.strftime("%d.%m.%Y")

    title = str(cfg.get("title") or "Звіт про переміщення ворога")
    subtitle_template = str(cfg.get("subtitle_template") or "за результатами радіорозвідки ({date})")
    subtitle = subtitle_template.format(date=date_str)

    group_header_template = str(cfg.get("group_header_template") or "В зоні функціонування {group} виявлено наступні переміщення:")
    line_template = str(cfg.get("line_template") or "{move} (р/м:{freq}, {unit}).")
    empty_text = str(cfg.get("empty_text") or "За наявними даними переміщень ворога не зафіксовано.")
    filename_template = str(cfg.get("filename_template") or "Переміщення ворога ({date}).docx")
    filename = filename_template.format(date=date_str)

    doc = Document()
    _set_base_font_12(doc)
    _add_header(doc, title=title, subtitle=subtitle)

    clean_items = [
        EnemyMoveItem(
            freq=(it.freq or "").strip(),
            move=(it.move or "").strip(),
            unit=(it.unit or "").strip(),
            group=(it.group or "").strip(),
        )
        for it in items
        if (it.move or "").strip() and (it.group or "").strip()
    ]

    if not clean_items:
        doc.add_paragraph(empty_text)
    else:
        grouped = _group_items(clean_items)
        for gi, group_name in enumerate(sorted(grouped.keys())):
            if gi != 0:
                doc.add_paragraph("")

            header_text = group_header_template.format(group=group_name)
            p = doc.add_paragraph()
            run = p.add_run(header_text)
            run.bold = True

            group_list = grouped[group_name]
            # Keep original order within each group (DB order).

            for it in group_list:
                line = line_template.format(move=it.move, freq=it.freq, unit=it.unit)
                doc.add_paragraph(line, style="List Bullet")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue(), filename


def today_sql_range() -> tuple[str, str]:
    """Return start/end ISO strings for today's local day."""
    d = datetime.now().date()
    start = datetime.combine(d, time.min).isoformat(timespec="seconds")
    end = datetime.combine(d, time.max.replace(microsecond=0)).isoformat(timespec="seconds")
    return start, end


__all__ = ["EnemyMoveItem", "build_enemy_moves_docx_bytes", "today_sql_range"]

