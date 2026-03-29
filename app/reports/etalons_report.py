from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.etalon_defaults import STANDARD_ETALON_OPERATION_MODE
from app.reports.enemy_moves_config import load_reports_config
from app.services.network_search import search_network_rows

DEFAULT_STATUS_IDS: list[int] = [1, 3, 7, 13, 14]


@dataclass(frozen=True)
class EtalonReportNetwork:
    network_id: int
    frequency: str
    unit: str
    zone: str
    group_name: str
    status_name: str
    callsigns: list[str]
    has_etalon: bool
    start_date: str | None
    end_date: str | None
    correspondents: str | None
    conclusions_text: str  # comma-separated from active network_tags.conclusions


def _set_base_font_12(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.size = Pt(12)


def _add_header(doc: Document, *, title: str, subtitle: str) -> None:
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)

    doc.add_paragraph("")


def _to_display_frequency(frequency: str, mask: str | None) -> str:
    freq = (frequency or "").strip()
    m = (mask or "").strip() if mask else ""
    return f"{freq} / {m}" if m else freq


def _in_placeholders(count: int) -> str:
    return ",".join(["?"] * count)


def _parse_date_any(value: str | None) -> date | None:
    """Parse DB TEXT date stored as ISO (YYYY-MM-DD) or as dd.mm.yyyy."""
    if not value:
        return None
    s = str(value).strip()
    if not s:
        return None
    # Most common: ISO date from HTML <input type="date">.
    try:
        return date.fromisoformat(s)
    except Exception:
        pass
    # Fallback: dd.mm.yyyy.
    for fmt in ("%d.%m.%Y", "%d.%m.%y"):
        try:
            return datetime.strptime(s, fmt).date()
        except Exception:
            continue
    return None


def _format_ddmmyyyy(d: date | None) -> str:
    if not d:
        return ""
    return d.strftime("%d.%m.%Y")


def _load_networks_for_report(conn, *, query: str, status_ids: list[int]) -> list[dict[str, Any]]:
    """Load networks for etalons report filtered by included status IDs.

    If query is provided - try to match networks by it (like the UI search).
    Otherwise - load all networks whose status_id is in the given list.
    """
    if not status_ids:
        status_ids = DEFAULT_STATUS_IDS

    q = (query or "").strip()
    placeholders_s = _in_placeholders(len(status_ids))

    if q:
        candidates = search_network_rows(conn, q, limit=500)
        ids = [int(r["id"]) for r in candidates if r and r["id"] is not None]
        if not ids:
            return []
        placeholders_n = _in_placeholders(len(ids))
        sql = f"""
            SELECT
              n.id,
              n.frequency,
              n.mask,
              n.unit,
              n.zone,
              COALESCE(g.name, '') AS group_name,
              n.status_id,
              COALESCE(s.name, '') AS status_name
            FROM networks n
            LEFT JOIN statuses s ON s.id = n.status_id
            LEFT JOIN groups g ON g.id = n.group_id
            WHERE n.id IN ({placeholders_n})
              AND n.status_id IN ({placeholders_s})
            ORDER BY n.frequency ASC, n.id ASC
        """
        return conn.execute(sql, [*ids, *status_ids]).fetchall()

    sql = f"""
        SELECT
          n.id,
          n.frequency,
          n.mask,
          n.unit,
          n.zone,
          COALESCE(g.name, '') AS group_name,
          n.status_id,
          COALESCE(s.name, '') AS status_name
        FROM networks n
        LEFT JOIN statuses s ON s.id = n.status_id
        LEFT JOIN groups g ON g.id = n.group_id
        WHERE n.status_id IN ({placeholders_s})
        ORDER BY n.frequency ASC, n.id ASC
    """
    return conn.execute(sql, status_ids).fetchall()


def _load_etalons_map(conn, *, network_ids: list[int]) -> dict[int, dict[str, Any]]:
    if not network_ids:
        return {}
    placeholders = _in_placeholders(len(network_ids))
    sql = f"""
        SELECT
          network_id,
          start_date,
          end_date,
          correspondents
        FROM etalons
        WHERE network_id IN ({placeholders})
    """
    rows = conn.execute(sql, network_ids).fetchall()
    out: dict[int, dict[str, Any]] = {}
    for r in rows:
        out[int(r["network_id"])] = r
    return out


def _load_callsigns_map(conn, *, network_ids: list[int]) -> dict[int, list[str]]:
    if not network_ids:
        return {}
    placeholders = _in_placeholders(len(network_ids))
    sql = f"""
        SELECT network_id, name
        FROM callsigns
        WHERE network_id IN ({placeholders})
        ORDER BY network_id ASC, name ASC
    """
    rows = conn.execute(sql, network_ids).fetchall()
    out: dict[int, list[str]] = {}
    for r in rows:
        nid = int(r["network_id"])
        out.setdefault(nid, []).append((r["name"] or "").strip())
    return out


def _load_conclusions_map(conn, *, network_ids: list[int]) -> dict[int, str]:
    """Return comma-separated conclusions from active network_tags per network.

    Only tags that have a non-empty ``conclusions`` field are included.
    Tags are ordered alphabetically by name for deterministic output.
    """
    if not network_ids:
        return {}
    placeholders = _in_placeholders(len(network_ids))
    sql = f"""
        SELECT ntl.network_id, nt.conclusions
        FROM network_tag_links ntl
        JOIN network_tags nt ON nt.id = ntl.tag_id
        WHERE ntl.network_id IN ({placeholders})
          AND nt.conclusions IS NOT NULL
          AND trim(nt.conclusions) != ''
        ORDER BY ntl.network_id ASC, nt.name ASC
    """
    rows = conn.execute(sql, network_ids).fetchall()
    out: dict[int, list[str]] = {}
    for r in rows:
        nid = int(r["network_id"])
        out.setdefault(nid, []).append((r["conclusions"] or "").strip())
    return {nid: ", ".join(parts) for nid, parts in out.items()}


def build_etalons_docx_bytes(
    *,
    conn,
    query: str = "",
    status_ids: list[int] | None = None,
    report_date: date | None = None,
) -> tuple[bytes, str]:
    cfg_root = load_reports_config()
    cfg = cfg_root.get("etalons") if isinstance(cfg_root, dict) else None
    cfg = cfg if isinstance(cfg, dict) else {}

    today = report_date or datetime.now().date()
    date_str = today.strftime("%d.%m.%Y")

    title = str(cfg.get("title") or "Звіт по еталонних описах")
    subtitle_template = str(cfg.get("subtitle_template") or "за результатами стану мереж ({date})")
    subtitle = subtitle_template.format(date=date_str)
    empty_text = str(cfg.get("empty_text") or "За наявними даними мереж для звіту не знайдено.")
    filename_template = str(cfg.get("filename_template") or "Еталонні описи ({date}).docx")
    filename = filename_template.format(date=date_str)

    included_status_ids = status_ids if status_ids else DEFAULT_STATUS_IDS

    networks = _load_networks_for_report(conn, query=query, status_ids=included_status_ids)
    network_ids = [int(n["id"]) for n in networks if n and n["id"] is not None]
    et_map = _load_etalons_map(conn, network_ids=network_ids)
    callsigns_map = _load_callsigns_map(conn, network_ids=network_ids)
    conclusions_map = _load_conclusions_map(conn, network_ids=network_ids)

    items: list[EtalonReportNetwork] = []
    for n in networks:
        nid = int(n["id"])
        et = et_map.get(nid)

        items.append(
            EtalonReportNetwork(
                network_id=nid,
                frequency=(n["frequency"] or "").strip(),
                unit=(n["unit"] or "").strip(),
                zone=(n["zone"] or "").strip(),
                group_name=(n["group_name"] or "").strip(),
                status_name=(n["status_name"] or "").strip(),
                callsigns=[c for c in callsigns_map.get(nid, []) if (c or "").strip()],
                has_etalon=bool(et),
                start_date=(et["start_date"] if et else None),
                end_date=(et["end_date"] if et else None),
                correspondents=(et["correspondents"] if et else None),
                conclusions_text=conclusions_map.get(nid, ""),
            )
        )

    doc = Document()
    _set_base_font_12(doc)
    _add_header(doc, title=title, subtitle=subtitle)

    if not items:
        doc.add_paragraph(empty_text)
    else:
        section_subtitle = str(
            cfg.get(
                "section_subtitle"
            )
            or "Для КХ та УКХ діапазонів радіохвиль (з урахуванням виявлених джерел маневреними групами РЕР частин РЕР у районах виконання завдань)"
        ).strip()
        modulation = str(cfg.get("modulation") or "NFM").strip()

        for idx, it in enumerate(items):
            doc.add_paragraph(section_subtitle)

            if not it.has_etalon:
                missing_text = f"Немає еталонного опису для частоти {it.frequency}"
                if it.status_name:
                    missing_text += f" (статус: {it.status_name})."
                doc.add_paragraph(missing_text)
            else:
                # Fixed form layout "Форма 1.5.3" — 9 points.
                net_name = f"УКХ р/м {it.unit}, н.п. {it.zone}".replace("  ", " ").strip()
                sd = _parse_date_any(it.start_date)
                start_part = _format_ddmmyyyy(sd) if sd else "невідомої дати"

                ed = _parse_date_any(it.end_date) or today
                end_part = _format_ddmmyyyy(ed) or date.today().strftime("%d.%m.%Y")

                # Point 3: auto-generated from network_tags.conclusions.
                if it.conclusions_text:
                    purpose_text = f"керування {it.conclusions_text}."
                else:
                    purpose_text = "—"

                doc.add_paragraph(f"1. Назва радіомережі: {net_name}")
                doc.add_paragraph(f"2. Район функціонування/розгортання: {it.zone}")
                doc.add_paragraph(f"3. Призначення: {purpose_text}")
                doc.add_paragraph(f"4. Склад кореспондентів: {(it.correspondents or '—').strip()}")
                doc.add_paragraph(
                    f"5. Позивні: {', '.join(it.callsigns) if it.callsigns else '—'}"
                )
                doc.add_paragraph(f"6. Частоти: {it.frequency}")
                if it.conclusions_text:
                    traffic_text = f"Службовий радіообмін для управління {it.conclusions_text}."
                else:
                    traffic_text = "Службовий радіообмін."

                doc.add_paragraph(f"7. Вид передачі: {modulation}")
                doc.add_paragraph(f"8. Режим роботи: {STANDARD_ETALON_OPERATION_MODE}")
                doc.add_paragraph(f"9. Характер роботи: {traffic_text}")
                doc.add_paragraph(f"10. Період функціонування: з {start_part} по {end_part}")

            if idx < len(items) - 1:
                doc.add_page_break()

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue(), filename


__all__ = ["build_etalons_docx_bytes"]
