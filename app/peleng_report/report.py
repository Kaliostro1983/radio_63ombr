"""DOCX report builder for peleng records (Form 1.2.13).

This module renders a list of peleng records into a DOCX document using
`python-docx`.

Usage in the system:
    - Web routes/services build a list of records (freq/unit_desc/dt/mgrs)
      and call `build_docx` to produce a DOCX file.
    - The runner script uses the same function for offline generation.

The report structure follows a fixed template: header, body tables
describing posts, and a final table listing peleng results.
"""

# src/pelengreport/report.py
# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable, Mapping, Sequence, Dict, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

from .config import load_config


_CFG = load_config()


def _add_header(doc: Document) -> None:
    """Add the fixed document header (title and date line)."""
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    normal.font.name = "Times New Roman"

    p = doc.add_paragraph(_CFG["header_form_label"])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")

    p = doc.add_paragraph(_CFG["header_title"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_str = datetime.now().strftime("%d.%m.%Y")
    sub_text = str(_CFG["header_sub_template"]).format(date=date_str)
    p = doc.add_paragraph(sub_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")


def _set_cell(cell, text: str, bold: bool = False,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              valign=WD_ALIGN_VERTICAL.CENTER):
    """Set table cell text and formatting with optional bold/alignment."""
    cell.text = ""
    lines = (text or "").split("\n")
    if not lines:
        lines = [""]

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
            p.text = line
        else:
            p = cell.add_paragraph(line)
        p.alignment = align
        if p.runs:
            p.runs[0].bold = bold
    cell.vertical_alignment = valign


def _post_label(post: Dict[str, Any]) -> str:
    """Build a display label for a post row (name + optional BP number)."""
    # name очікуємо типу: "МАЯКИ"
    # bp_number: "0001"
    name = (post.get("name") or "").strip()
    bp = (post.get("bp_number") or "").strip()
    if bp:
        return f"{name},\nБП №{bp}"
    return name


def _add_body(doc: Document, total_pelengs: int, posts: Sequence[Dict[str, Any]]) -> None:
    """Add the body sections and tables describing posts and totals."""
    # 1.
    p = doc.add_paragraph(str(_CFG["section1_text"]))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tbl1 = doc.add_table(rows=1, cols=5)
    tbl1.style = "Table Grid"
    hdrs1 = list(_CFG["table1_headers"])
    for j, h in enumerate(hdrs1):
        _set_cell(tbl1.rows[0].cells[j], h, bold=True)

    row = tbl1.add_row()
    row.cells[2].merge(row.cells[3]).merge(row.cells[4])
    _set_cell(row.cells[2], str(_CFG["table1_group_label"]), bold=True)
    _set_cell(row.cells[0], "")
    _set_cell(row.cells[1], "")

    for idx, post in enumerate(posts, 1):
        row = tbl1.add_row()
        _set_cell(row.cells[0], f"{idx}.")
        _set_cell(row.cells[1], post.get("unit", _CFG["default_unit"]))
        _set_cell(row.cells[2], _post_label(post), align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row.cells[3], post.get("equipment", _CFG["default_equipment"]))
        _set_cell(
            row.cells[4],
            post.get("task_progress", _CFG["default_task_progress"]),
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

    doc.add_paragraph("")

    # 2.
    p = doc.add_paragraph(str(_CFG["section2_text"]))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(str(_CFG["section2_no_changes_text"]))
    if p.runs:
        p.runs[0].italic = True

    # 3.
    p = doc.add_paragraph(str(_CFG["section3_text"]))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tbl2 = doc.add_table(rows=1, cols=6)
    tbl2.style = "Table Grid"
    hdrs2 = list(_CFG["table2_headers"])
    for j, h in enumerate(hdrs2):
        _set_cell(tbl2.rows[0].cells[j], h, bold=True)

    row = tbl2.add_row()
    row.cells[2].merge(row.cells[3]).merge(row.cells[4])
    _set_cell(row.cells[2], str(_CFG["table2_group_label"]), bold=True)
    _set_cell(row.cells[0], "")
    _set_cell(row.cells[1], "")
    _set_cell(row.cells[5], "")

    for idx, post in enumerate(posts, 1):
        row = tbl2.add_row()
        _set_cell(row.cells[0], f"{idx}.")
        _set_cell(row.cells[1], post.get("unit", _CFG["default_unit"]))
        _set_cell(row.cells[2], _post_label(post), align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row.cells[3], post.get("equipment", _CFG["default_equipment"]))
        _set_cell(row.cells[4], str(total_pelengs if idx == 1 else 0))
        _set_cell(row.cells[5], post.get("note", ""))

    doc.add_paragraph("")
    p = doc.add_paragraph(str(_CFG["section4_text"]))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")


def _add_table(doc: Document, rows: Iterable[Mapping[str, str]]) -> None:
    """Add the main results table listing peleng records."""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    
    from docx.shared import Cm

    table.autofit = False

    # Базова ширина
    total_width = Cm(16)

    # Збільшуємо 5 колонку на 30%
    col1 = Cm(2.0)
    col2 = Cm(2.5)
    col3 = Cm(4.5)
    col4 = Cm(3.0)
    col5 = Cm(4.0)  # було ~3 → стало більше

    widths = [col1, col2, col3, col4, col5]

    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width
    
    hdr = table.rows[0].cells
    headers = list(_CFG["results_table_headers"])
    # Ensure we have exactly 5 headers; fall back to defaults length-wise.
    while len(headers) < 5:
        headers.append("")
    headers = headers[:5]
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = headers

    for i, rec in enumerate(rows, 1):
        cells = table.add_row().cells
        _set_cell(cells[0], str(i))  # центр
        _set_cell(cells[1], str(rec.get("freq_or_mask", "")))  # центр
        _set_cell(cells[2], str(rec.get("unit_desc", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(cells[3], str(rec.get("dt", "")))  # центр
        _set_cell(cells[4], str(rec.get("mgrs", "")))  # центр


def build_docx(records: Sequence[Mapping[str, str]], out_path: Path, posts: Sequence[Dict[str, Any]]) -> Path:
    """Build and save a DOCX report to disk.

    Args:
        records: iterable of record mappings with keys `freq_or_mask`,
            `unit_desc`, `dt`, `mgrs`.
        out_path: output file path.
        posts: list of post dicts used to populate header tables.

    Returns:
        Path: the output path (same as `out_path`).
    """
    doc = Document()
    _add_header(doc)
    _add_body(doc, total_pelengs=len(records), posts=posts)
    _add_table(doc, records)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


__all__ = ["build_docx"]